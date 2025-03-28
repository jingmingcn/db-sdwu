from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# 创建一个新的Word文档
doc = Document()

# 添加封面
doc.add_page_break()
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 学校名称
school_name = cover.add_run("山东女子学院\n")
school_name.bold = True
school_name.font.size = Pt(18)

# 学院名称
college_name = cover.add_run("数据科学与计算机学院\n")
college_name.font.size = Pt(16)

# 课程名称
course_name = cover.add_run("《数据库系统概论》作业\n\n")
course_name.bold = True
course_name.font.size = Pt(18)

# 添加学生信息部分
student_info = doc.add_paragraph()
student_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
student_info.add_run("学号：______________________\n").font.size = Pt(14)
student_info.add_run("姓名：______________________\n").font.size = Pt(14)
student_info.add_run("专业：______________________\n").font.size = Pt(14)
student_info.add_run("班级：______________________\n").font.size = Pt(14)

# 添加分隔符
doc.add_page_break()

# 添加作业要求部分
doc.add_heading("作业要求", level=1)
doc.add_paragraph("1. 请认真阅读并完成以下题目。\n2. 作业需独立完成，严禁抄袭。\n3. 提交格式：电子版或手写扫描版。\n")

# 添加作业题目部分
doc.add_heading("作业题目", level=1)
doc.add_paragraph("1. 解释数据库管理系统（DBMS）的主要功能。\n")
doc.add_paragraph("2. 关系数据库的三大范式分别是什么？请举例说明。\n")
doc.add_paragraph("3. SQL 语言中的 DDL、DML 和 DCL 分别指什么？请举例说明。\n")
doc.add_paragraph("4. 设计一个简单的学生成绩管理数据库，包括至少三个表，并给出表结构定义（列名、数据类型、主键/外键等）。\n")

# 保存文件
file_path = "数据库系统概论作业模板.docx"
doc.save(file_path)

file_path